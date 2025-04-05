import os
import docx
import pandas as pd
import re
import argparse
from openpyxl import Workbook
import json
import requests
from typing import Dict, List, Any

class LocalLLMInterface:
    """Base class for local LLM interfaces"""
    def generate(self, prompt: str) -> str:
        raise NotImplementedError("Subclasses must implement generate()")

class OllamaInterface(LocalLLMInterface):
    """Interface for Ollama LLMs"""
    def __init__(self, model: str = "llama3"):
        self.model = model
        self.api_base = "http://localhost:11434/api"
        # Check if Ollama is running
        try:
            requests.get(f"{self.api_base}/version")
        except requests.exceptions.ConnectionError:
            raise ConnectionError("Ollama server not running. Start with 'ollama serve'")

    def generate(self, prompt: str) -> str:
        """Generate text using Ollama API"""
        response = requests.post(
            f"{self.api_base}/generate",
            json={"model": self.model, "prompt": prompt, "stream": False}
        )
        if response.status_code == 200:
            return response.json().get("response", "")
        else:
            raise Exception(f"Ollama API error: {response.status_code} - {response.text}")

class LMStudioInterface(LocalLLMInterface):
    """Interface for LM Studio"""
    def __init__(self, port: int = 1234):
        self.api_base = f"http://localhost:{port}/v1"
        # Check if LM Studio is running
        try:
            requests.get(f"{self.api_base}/models")
        except requests.exceptions.ConnectionError:
            raise ConnectionError("LM Studio not running. Start LM Studio and enable API server.")

    def generate(self, prompt: str) -> str:
        """Generate text using LM Studio API"""
        response = requests.post(
            f"{self.api_base}/chat/completions",
            json={
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.2
            }
        )
        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"]
        else:
            raise Exception(f"LM Studio API error: {response.status_code} - {response.text}")

class TextGenerationWebUIInterface(LocalLLMInterface):
    """Interface for Text Generation Web UI"""
    def __init__(self, port: int = 5000):
        self.api_base = f"http://localhost:{port}/api"
        # Check if Text Generation Web UI is running
        try:
            requests.get(f"{self.api_base}/v1/models")
        except requests.exceptions.ConnectionError:
            raise ConnectionError("Text Generation Web UI not running. Start the server first.")

    def generate(self, prompt: str) -> str:
        """Generate text using Text Generation Web UI API"""
        response = requests.post(
            f"{self.api_base}/v1/generate",
            json={
                "prompt": prompt,
                "max_new_tokens": 1024,
                "temperature": 0.2
            }
        )
        if response.status_code == 200:
            return response.json().get("results", [{}])[0].get("text", "")
        else:
            raise Exception(f"Text Generation Web UI API error: {response.status_code} - {response.text}")

class WordToExcelConverter:
    def __init__(self, llm_interface: LocalLLMInterface = None, llm_type: str = "ollama", model: str = "llama3"):
        """
        Initialize the converter with a local LLM interface
        
        Args:
            llm_interface: Custom LLM interface (optional)
            llm_type: Type of LLM interface to use if not provided (ollama, lmstudio, textgen)
            model: Model name for Ollama
        """
        if llm_interface:
            self.llm = llm_interface
        else:
            if llm_type == "ollama":
                self.llm = OllamaInterface(model=model)
            elif llm_type == "lmstudio":
                self.llm = LMStudioInterface()
            elif llm_type == "textgen":
                self.llm = TextGenerationWebUIInterface()
            else:
                raise ValueError(f"Unknown LLM type: {llm_type}")

    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract full text from a Word document"""
        doc = docx.Document(docx_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                if any(text for text in row_text):
                    full_text.append(" | ".join(row_text))
                
        return "\n".join(full_text)

    def analyze_content(self, text: str) -> Dict[str, Any]:
        """Use LLM to analyze document content and suggest table structure"""
        prompt = f"""You are a document analysis expert who specializes in extracting structured data.
        
        Analyze this document content and determine how to structure it as spreadsheet data.
        Identify tables, lists, key-value pairs, or any structured data that should be extracted.
        
        Return your response as a JSON object with the following structure:
        {{
            "tables": [
                {{
                    "name": "Table name or description",
                    "columns": ["Column1", "Column2", ...],
                    "extraction_rules": "Explain how to identify and extract this data"
                }}
            ],
            "analysis": "Brief analysis of the document structure and content"
        }}
        
        Document content:
        {text[:2000]}  # Limiting to first 2000 chars to avoid context window issues
        
        Respond ONLY with the JSON object.
        """
        
        result = self.llm.generate(prompt)
            
        # Extract JSON from response
        try:
            # Find JSON structure in the response
            json_match = re.search(r'({[\s\S]*})', result)
            if json_match:
                json_str = json_match.group(1)
                return json.loads(json_str)
            else:
                # If no JSON found, try to parse the entire response
                return json.loads(result)
        except json.JSONDecodeError:
            print("Error parsing LLM response. Using fallback method.")
            # Fallback to simple table extraction
            return {
                "tables": [{
                    "name": "ExtractedData",
                    "columns": ["Content"],
                    "extraction_rules": "Extract all content as raw text"
                }],
                "analysis": "Could not analyze document structure. Using raw text extraction."
            }

    def extract_structured_data(self, text: str, table_spec: Dict[str, Any]) -> List[Dict[str, str]]:
        """Extract structured data based on LLM's analysis"""
        prompt = f"""You are a data extraction expert who excels at structuring information from documents.
        
        Extract data from this document according to these instructions:
        Table name: {table_spec['name']}
        Columns: {', '.join(table_spec['columns'])}
        Extraction rules: {table_spec['extraction_rules']}
        
        Return ONLY a JSON array of objects where keys are column names.
        Format as: [{{column1: value, column2: value}}, {{column1: value, column2: value}}]
        
        Document content:
        {text}
        """
        
        result = self.llm.generate(prompt)
            
        # Extract JSON array from response
        try:
            # Find JSON structure in the response
            json_match = re.search(r'(\[[\s\S]*\])', result)
            if json_match:
                json_str = json_match.group(1)
                return json.loads(json_str)
            else:
                # If no JSON found, try to parse the entire response
                return json.loads(result)
        except json.JSONDecodeError:
            print(f"Error parsing extracted data for {table_spec['name']}. Using fallback.")
            # Return dummy data with column names
            return [{col: f"Error extracting {col}" for col in table_spec['columns']}]

    def chunked_analysis(self, text: str) -> Dict[str, Any]:
        """Handle large documents by analyzing in chunks"""
        # Split document into chunks of approximately 2000 characters
        chunk_size = 2000
        chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
        
        # Analyze first chunk to get initial structure
        analysis = self.analyze_content(chunks[0])
        
        # If document is small enough, return the analysis
        if len(chunks) == 1:
            return analysis
            
        # For larger documents, update the analysis with additional chunks
        for i, chunk in enumerate(chunks[1:], start=1):
            prompt = f"""You previously analyzed part of a document and found these potential data tables:
            {json.dumps(analysis['tables'])}
            
            Here is part {i+1} of the document. Update your analysis if you find new information:
            {chunk}
            
            Return the updated analysis as a JSON object with the same structure as before.
            """
            
            try:
                result = self.llm.generate(prompt)
                update = json.loads(re.search(r'({[\s\S]*})', result).group(1))
                
                # Merge in any new tables
                existing_names = [t['name'] for t in analysis['tables']]
                for table in update.get('tables', []):
                    if table['name'] not in existing_names:
                        analysis['tables'].append(table)
                
                # Update analysis text
                analysis['analysis'] += " " + update.get('analysis', '')
                
            except (json.JSONDecodeError, AttributeError):
                # If parsing fails, continue with current analysis
                pass
                
        return analysis

    def convert_to_excel(self, word_path: str, excel_path: str = None) -> str:
        """Convert Word document to Excel workbook"""
        if not excel_path:
            # Create output path based on input filename
            excel_path = os.path.splitext(word_path)[0] + '.xlsx'
            
        # Extract text from Word document
        text = self.extract_text_from_docx(word_path)
        
        # Analyze content to determine structure
        analysis = self.chunked_analysis(text)
        
        print(f"Analysis complete. Found {len(analysis['tables'])} potential tables.")
        
        # Create Excel workbook
        workbook = Workbook()
        # Remove default sheet
        default_sheet = workbook.active
        workbook.remove(default_sheet)
        
        # For each table found by the model, create a sheet and add data
        for table_spec in analysis['tables']:
            sheet = workbook.create_sheet(title=table_spec['name'][:31])  # Excel sheet name limit 31 chars
            sheet.append(table_spec['columns'])  # Add column headers
            
            # Extract data for this table
            for row_data in self.extract_structured_data(text, table_spec):
                sheet.append([row_data.get(col, "") for col in table_spec['columns']])
        
        # Save the workbook
        workbook.save(excel_path)
        return excel_path

def main():
    parser = argparse.ArgumentParser(description="Convert Word to Excel with structured data.")
    parser.add_argument("word_path", type=str, help="Path to the Word document to convert")
    parser.add_argument("--excel_path", type=str, help="Optional path to save the Excel file")
    parser.add_argument("--llm_type", type=str, default="ollama", choices=["ollama", "lmstudio", "textgen"],
                        help="Specify the type of LLM to use")
    parser.add_argument("--model", type=str, default="llama3", help="Model to use with Ollama")
    
    args = parser.parse_args()
    
    # Set up the LLM interface
    converter = WordToExcelConverter(
        llm_type=args.llm_type,
        model=args.model
    )
    
    # Convert Word document to Excel
    excel_path = converter.convert_to_excel(args.word_path, args.excel_path)
    print(f"Conversion complete. Saved as: {excel_path}")
    return excel_path

if __name__ == "__main__":
    main()